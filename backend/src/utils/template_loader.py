"""
模板加载工具模块

本模块负责加载和解析.docx格式的回复模板文件，包括：
- Interim reply 模板
- Final reply 模板
- Wrong referral reply 模板

主要功能：
1. 读取.docx文件并提取文本内容
2. 缓存模板内容以提高性能
3. 解析模板中的多个示例

作者: Project3 Team
版本: 1.0
"""
import os
from docx import Document
from typing import Dict, List, Optional


class TemplateLoader:
    """模板加载器类"""
    
    def __init__(self, templates_dir: str = None):
        """
        初始化模板加载器
        
        Args:
            templates_dir: 模板文件目录路径，默认为项目根目录下的docs/templates
        """
        if templates_dir is None:
            # 获取项目根目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(current_dir)))
            templates_dir = os.path.join(project_root, "docs", "templates")
        
        self.templates_dir = templates_dir
        self.templates_cache: Dict[str, str] = {}
        
        # 模板文件映射
        self.template_files = {
            "interim": "Interim reply  Anwser and Template.docx",
            "final": "Final reply answer and Template.docx",
            "wrong_referral": "Wrong referral reply Answer and template.docx"
        }
    
    def load_template(self, reply_type: str) -> Optional[str]:
        """
        加载并缓存模板内容
        
        Args:
            reply_type: 回复类型 (interim, final, wrong_referral)
        
        Returns:
            模板内容字符串，失败返回None
        """
        # 检查缓存
        if reply_type in self.templates_cache:
            return self.templates_cache[reply_type]
        
        # 检查reply_type是否有效
        if reply_type not in self.template_files:
            print(f"❌ 无效的回复类型: {reply_type}")
            return None
        
        # 构建文件路径
        template_file = self.template_files[reply_type]
        file_path = os.path.join(self.templates_dir, template_file)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"❌ 模板文件不存在: {file_path}")
            return None
        
        try:
            # 读取.docx文件
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 只保留非空段落
                    paragraphs.append(text)
            
            # 提取表格内容（如果有）
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            paragraphs.append(" | ".join(row_data))
            
            # 合并所有内容
            template_content = "\n\n".join(paragraphs)
            
            # 缓存模板内容
            self.templates_cache[reply_type] = template_content
            
            print(f"✅ 成功加载模板: {reply_type} ({len(template_content)} 字符)")
            return template_content
            
        except Exception as e:
            print(f"❌ 加载模板文件失败: {e}")
            return None
    
    def parse_template_examples(self, template_content: str) -> List[str]:
        """
        解析模板中的多个示例
        
        Args:
            template_content: 模板内容字符串
        
        Returns:
            示例列表
        """
        if not template_content:
            return []
        
        # 按照空行分割示例
        examples = []
        current_example = []
        
        for line in template_content.split('\n'):
            line = line.strip()
            if line:
                current_example.append(line)
            else:
                if current_example:
                    examples.append('\n'.join(current_example))
                    current_example = []
        
        # 添加最后一个示例
        if current_example:
            examples.append('\n'.join(current_example))
        
        return examples
    
    def get_all_templates(self) -> Dict[str, str]:
        """
        获取所有模板内容
        
        Returns:
            包含所有模板的字典 {reply_type: template_content}
        """
        all_templates = {}
        for reply_type in self.template_files.keys():
            content = self.load_template(reply_type)
            if content:
                all_templates[reply_type] = content
        return all_templates
    
    def clear_cache(self):
        """清除模板缓存"""
        self.templates_cache.clear()
        print("✅ 模板缓存已清除")


# 全局模板加载器实例
_template_loader_instance = None


def get_template_loader() -> TemplateLoader:
    """
    获取全局模板加载器实例（单例模式）
    
    Returns:
        TemplateLoader实例
    """
    global _template_loader_instance
    if _template_loader_instance is None:
        _template_loader_instance = TemplateLoader()
    return _template_loader_instance
