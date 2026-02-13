#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG文件存储管理模块

本模块提供知识库文件的持久化存储、管理和预览功能。

主要功能：
1. 保存上传的RAG文件到指定目录
2. 删除文件及其元数据
3. 生成文件预览内容
4. 文件路径管理

存储结构：
backend/data/rag_files/{year}/{month}/{filename}

作者: Project3 Team
版本: 1.0
"""

import os
import shutil
import re
from datetime import datetime
from typing import Optional, Tuple
import pytz
from pathlib import Path

# Default short preview length (stored in DB)
DEFAULT_PREVIEW_LENGTH = 500


def sanitize_filename(original_filename: str, max_base_length: int = 120) -> str:
    """
    Sanitize user-provided filename to prevent traversal and unsafe characters.
    """
    name = Path(original_filename or "").name.replace("\x00", "").strip()
    if not name:
        return "upload.bin"
    base, ext = os.path.splitext(name)
    safe_base = re.sub(r"[^A-Za-z0-9._-]", "_", base).strip("._")
    safe_base = safe_base[:max_base_length] or "upload"
    safe_ext = re.sub(r"[^A-Za-z0-9.]", "", ext)[:10]
    return f"{safe_base}{safe_ext}"


def get_rag_storage_path() -> str:
    """
    获取RAG文件存储根目录
    
    Returns:
        str: RAG文件存储根目录的绝对路径
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    backend_dir = os.path.dirname(os.path.dirname(current_dir))
    storage_path = os.path.join(backend_dir, 'data', 'rag_files')
    
    # 确保目录存在
    os.makedirs(storage_path, exist_ok=True)
    
    return storage_path


def save_rag_file(file_content: bytes, original_filename: str) -> Tuple[str, str]:
    """
    保存RAG文件到存储目录
    
    Args:
        file_content: 文件二进制内容
        original_filename: 原始文件名
        
    Returns:
        Tuple[str, str]: (完整文件路径, 相对路径)
        
    Example:
        full_path, relative_path = save_rag_file(file_content, "知识库.xlsx")
        # full_path: /path/to/backend/data/rag_files/2024/02/知识库_1707123456.xlsx
        # relative_path: 2024/02/知识库_1707123456.xlsx
    """
    # 获取北京时间
    beijing_tz = pytz.timezone('Asia/Shanghai')
    now = datetime.now(beijing_tz)
    
    # 创建年月目录
    year = now.strftime('%Y')
    month = now.strftime('%m')
    
    storage_root = get_rag_storage_path()
    storage_dir = os.path.join(storage_root, year, month)
    os.makedirs(storage_dir, exist_ok=True)
    
    # 生成唯一文件名（添加时间戳避免重复）
    timestamp = int(now.timestamp())
    safe_original_filename = sanitize_filename(original_filename)
    filename_parts = os.path.splitext(safe_original_filename)
    unique_filename = f"{filename_parts[0]}_{timestamp}{filename_parts[1]}"
    
    # 保存文件
    full_path = os.path.join(storage_dir, unique_filename)
    with open(full_path, 'wb') as f:
        f.write(file_content)
    
    # 计算相对路径
    relative_path = os.path.join(year, month, unique_filename)
    
    print(f"✅ 文件保存成功: {relative_path}")
    return full_path, relative_path


def delete_rag_file(file_path: str) -> bool:
    """
    删除RAG文件
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对路径）
        
    Returns:
        bool: 删除成功返回True，失败返回False
    """
    try:
        # 如果是相对路径，转换为绝对路径
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        
        # 检查文件是否存在
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"✅ 文件删除成功: {file_path}")
            
            # 尝试删除空目录
            _clean_empty_dirs(os.path.dirname(file_path))
            return True
        else:
            print(f"⚠️ 文件不存在: {file_path}")
            return False
            
    except Exception as e:
        print(f"❌ 文件删除失败: {e}")
        return False


def _clean_empty_dirs(directory: str):
    """
    清理空目录（递归向上清理）
    
    Args:
        directory: 要检查的目录路径
    """
    try:
        storage_root = get_rag_storage_path()
        
        # 只清理storage_root下的目录
        if not directory.startswith(storage_root):
            return
        
        # 如果目录为空且不是根目录，则删除
        if os.path.exists(directory) and not os.listdir(directory) and directory != storage_root:
            os.rmdir(directory)
            print(f"🗑️ 清理空目录: {directory}")
            # 递归检查父目录
            _clean_empty_dirs(os.path.dirname(directory))
    except Exception as e:
        print(f"⚠️ 清理目录时出错: {e}")


def get_file_preview(file_path: str, file_type: str, max_length: int = 500) -> Optional[str]:
    """
    生成文件预览内容
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对路径）
        file_type: 文件类型（excel, word, pdf, txt, csv等）
        max_length: 最大预览字符数（默认500）
        
    Returns:
        Optional[str]: 预览文本，失败返回None
    """
    try:
        # 如果是相对路径，转换为绝对路径
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"⚠️ 文件不存在: {file_path}")
            return None
        
        preview_text = ""
        
        # 根据文件类型提取预览
        if file_type in ['txt', 'csv']:
            # 文本文件直接读取
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                preview_text = f.read(max_length)
                
        elif file_type == 'excel':
            # Excel文件预览
            try:
                import pandas as pd
                excel_file = pd.ExcelFile(file_path)
                # 读取第一个sheet的前几行
                df = pd.read_excel(file_path, sheet_name=0, nrows=10)
                preview_text = df.to_string()[:max_length]
            except Exception as e:
                print(f"⚠️ Excel预览失败: {e}")
                preview_text = "[Excel文件，无法生成预览]"
                
        elif file_type == 'word':
            # Word文件预览
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs[:5]]
                preview_text = "\n".join(paragraphs)[:max_length]
            except Exception as e:
                print(f"⚠️ Word预览失败: {e}")
                preview_text = "[Word文件，无法生成预览]"
                
        elif file_type == 'pdf':
            # PDF文件预览
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    if len(pdf.pages) > 0:
                        text = pdf.pages[0].extract_text()
                        preview_text = text[:max_length] if text else "[PDF文件，无文本内容]"
                    else:
                        preview_text = "[PDF文件为空]"
            except Exception as e:
                print(f"⚠️ PDF预览失败: {e}")
                preview_text = "[PDF文件，无法生成预览]"
                
        elif file_type == 'image':
            # 图片文件不生成文本预览
            preview_text = "[图片文件]"
            
        else:
            preview_text = f"[{file_type}文件，不支持预览]"
        
        return preview_text
        
    except Exception as e:
        print(f"❌ 生成预览失败: {e}")
        return None


def get_file_preview_slice(
    file_path: str,
    file_type: str,
    offset: int = 0,
    limit: Optional[int] = None,
) -> Tuple[Optional[str], int]:
    """
    获取文件预览内容（支持分页/全量）。
    
    Args:
        file_path: 文件路径（绝对或相对）
        file_type: 文件类型（excel, word, pdf, txt, csv等）
        offset: 起始字符偏移（默认0）
        limit: 返回最大字符数，None 表示不限制（全量）
        
    Returns:
        Tuple[Optional[str], int]: (预览文本, 文件总字符数)，失败时 (None, 0)
    """
    try:
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        
        if not os.path.exists(file_path):
            print(f"⚠️ 文件不存在: {file_path}")
            return None, 0
        
        full_text: Optional[str] = None
        
        if file_type in ['txt', 'csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
        elif file_type == 'excel':
            try:
                import pandas as pd
                df = pd.read_excel(file_path, sheet_name=0)
                full_text = df.to_string()
            except Exception as e:
                print(f"⚠️ Excel 全量读取失败: {e}")
                return "[Excel文件，无法生成全文预览]", 0
        elif file_type == 'word':
            try:
                from docx import Document
                doc = Document(file_path)
                full_text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                print(f"⚠️ Word 全量读取失败: {e}")
                return "[Word文件，无法生成全文预览]", 0
        elif file_type == 'pdf':
            try:
                import pdfplumber
                parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            parts.append(t)
                full_text = "\n\n".join(parts) if parts else "[PDF文件，无文本内容]"
            except Exception as e:
                print(f"⚠️ PDF 全量读取失败: {e}")
                return "[PDF文件，无法生成全文预览]", 0
        elif file_type == 'image':
            return "[图片文件]", 0
        else:
            return f"[{file_type}文件，不支持预览]", 0
        
        if full_text is None:
            return None, 0
        
        total = len(full_text)
        if limit is None:
            snippet = full_text[offset:] if offset > 0 else full_text
        else:
            end = min(offset + limit, total)
            snippet = full_text[offset:end]
        return snippet, total
        
    except Exception as e:
        print(f"❌ get_file_preview_slice 失败: {e}")
        return None, 0


def get_absolute_path(relative_path: str) -> str:
    """
    将相对路径转换为绝对路径
    
    Args:
        relative_path: 相对路径
        
    Returns:
        str: 绝对路径
    """
    if os.path.isabs(relative_path):
        return relative_path
    
    storage_root = get_rag_storage_path()
    return os.path.join(storage_root, relative_path)


def file_exists(file_path: str) -> bool:
    """
    检查文件是否存在
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对路径）
        
    Returns:
        bool: 文件存在返回True，否则返回False
    """
    if not os.path.isabs(file_path):
        file_path = get_absolute_path(file_path)
    
    return os.path.exists(file_path)


def get_file_size(file_path: str) -> Optional[int]:
    """
    获取文件大小
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对路径）
        
    Returns:
        Optional[int]: 文件大小（bytes），失败返回None
    """
    try:
        if not os.path.isabs(file_path):
            file_path = get_absolute_path(file_path)
        
        if os.path.exists(file_path):
            return os.path.getsize(file_path)
        return None
    except Exception as e:
        print(f"❌ 获取文件大小失败: {e}")
        return None
